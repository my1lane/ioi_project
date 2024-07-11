import os
from docx import Document
import lxml.etree as ET

current_dir = os.path.dirname(os.path.realpath(__file__))

file_path = os.path.join(current_dir, 'new_docx.docx')
output_file = os.path.join(current_dir, 'output_ver2.xml')


def get_style_info(run):
    font = run.font.name if run.font.name else "None"
    size = f"{run.font.size.pt} pt" if run.font.size else "None"
    bold = "Да" if run.bold else "Нет"
    return font, size, bold


def parse_paragraph(paragraph, is_toc=False):
    para_xml = ET.Element("toc_content" if is_toc else "paragraph")
    for run in paragraph.runs:
        font, size, bold = get_style_info(run)
        run_xml = ET.SubElement(para_xml, "content")
        run_xml.text = run.text
        run_xml.set("font", font)
        run_xml.set("size", size)
        run_xml.set("bold", bold)
    return para_xml


def parse_table(table):
    table_xml = ET.Element("table")
    for row in table.rows:
        row_xml = ET.SubElement(table_xml, "row")
        for cell in row.cells:
            cell_xml = ET.SubElement(row_xml, "cell")
            for paragraph in cell.paragraphs:
                cell_xml.append(parse_paragraph(paragraph))
    return table_xml


def parse_document(doc):
    root = ET.Element("document")
    is_toc_section = False

    for paragraph in doc.paragraphs:
        # Проверка на заголовок оглавления
        if paragraph.text.strip().lower() in ["оглавление", "содержание"]:
            is_toc_section = True
            continue

        # Если это оглавление, обрабатываем как toc_content
        if is_toc_section:
            root.append(parse_paragraph(paragraph, is_toc=True))
            is_toc_section = False
        else:
            root.append(parse_paragraph(paragraph))

    for table in doc.tables:
        root.append(parse_table(table))

    return root


def save_to_xml(root, filename):
    tree = ET.ElementTree(root)
    with open(filename, 'wb') as f:
        tree.write(f, encoding="utf-8", xml_declaration=True, pretty_print=True)


def main(docx_filename, xml_filename):
    doc = Document(docx_filename)
    root = parse_document(doc)
    save_to_xml(root, xml_filename)


if __name__ == "__main__":
    main(file_path, output_file)
