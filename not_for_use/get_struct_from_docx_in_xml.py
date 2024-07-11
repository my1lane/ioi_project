import os
from docx import Document
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET
from xml.dom import minidom

current_dir = os.path.dirname(os.path.realpath(__file__))

file_path = os.path.join(current_dir, 'new_docx.docx')

# Открываем документ
doc = Document(file_path)

# Функция для получения информации о стиле текста
def get_info_style(run, paragraph):
    # Попробуем получить информацию на уровне run
    font_name = run.font.name
    font_size = run.font.size.pt if run.font.size else None

    # Если информация недоступна, попробуем получить ее на уровне параграфа
    if not font_name:
        font_name = paragraph.style.font.name if paragraph.style.font.name else "Unknown"
    if not font_size:
        font_size = paragraph.style.font.size.pt if paragraph.style.font.size else "Unknown"

    bold = run.bold
    italic = run.italic

    info_style = {
        'font': font_name,
        'size': f"{font_size}pt" if font_size != "Unknown" else "Unknown",
        'bold': 'yes' if bold else 'no',
        'italic': 'yes' if italic else 'no'
    }

    return info_style


# Функция для создания XML элемента
def create_xml_element(element_type, text, info_style):
    element = ET.Element(element_type)
    element.text = text
    for key, value in info_style.items():
        if value:
            element.set(key, value)
    return element


# Функция для получения текста до первой точки
def get_text_before_dot(text):
    return text.split('.')[0] if '.' in text else text


root = ET.Element("document")

heading_counter = 1
table_counter = 1
toc_started = False
toc_ended = False

# Проверяем, начинается ли содержание
for para in doc.paragraphs:
    if 'Содержание' in para.text:
        toc_started = True
        break

# Проходим по всем элементам документа и добавляем их в XML
for para in doc.paragraphs:
    # Пропускаем пустые строки
    if not para.text.strip():
        continue

    # Обработка содержания
    if toc_started and not toc_ended:
        if para.style.name.startswith('TOC Heading') or 'Содержание' in para.style.name:
            element = create_xml_element("content", para.text, {})
            root.append(element)
            continue
        elif not para.style.name.startswith('TOC Heading') and not 'Содержание' in para.style.name:
            toc_ended = True

    # Получаем текст до первой точки для всех элементов, кроме содержания
    text_to_use = para.text if toc_started and not toc_ended else get_text_before_dot(para.text)

    # Получаем информацию о стиле первого run
    if para.runs:
        style_info = get_info_style(para.runs[0], para)
    else:
        style_info = {"font": "Unknown", "size": "Unknown", "bold": "no", "italic": "no"}

    # Вывод отладочной информации о стиле
    print(f"Параграф: {text_to_use}")
    print(f"Стиль: {style_info}")

    # Обработка заголовков, списков, рисунков и параграфов
    if para.style.name.startswith('Heading'):
        element = create_xml_element(f"heading{heading_counter}", text_to_use, style_info)
        heading_counter += 1
    elif para.style.name in ['List Bullet', 'List Number']:
        element = create_xml_element("list", text_to_use, style_info)
    elif any('drawing' in run._element.xml for run in para.runs):
        element = create_xml_element("figure", text_to_use, style_info)
    else:
        element = create_xml_element("paragraph", text_to_use, style_info)

    root.append(element)

# Проходим по всем таблицам в документе и добавляем их в XML
for table in doc.tables:
    table_element = ET.Element(f"table{table_counter}")
    for row in table.rows:
        row_element = ET.Element("row")
        for cell in row.cells:
            text_before_dot = get_text_before_dot(cell.text)
            if cell.paragraphs and cell.paragraphs[0].runs:
                style_info = get_info_style(cell.paragraphs[0].runs[0], cell.paragraphs[0])
            else:
                style_info = {"font": "Unknown", "size": "Unknown", "bold": "no", "italic": "no"}
            cell_element = create_xml_element("cell", text_before_dot, style_info)
            row_element.append(cell_element)
        table_element.append(row_element)
    root.append(table_element)
    table_counter += 1

# Сохраняем XML-документ в файл с отступами
output_file = os.path.join(current_dir, 'output.xml')
tree = ET.ElementTree(root)

# Форматируем XML с отступами
xml_str = ET.tostring(root, encoding='utf-8')
parsed_str = minidom.parseString(xml_str)
pretty_xml_as_str = parsed_str.toprettyxml(indent="  ")

with open(output_file, 'w', encoding='utf-8') as f:
    f.write(pretty_xml_as_str)

print(f"Документ успешно сконвертирован в XML и сохранен как {output_file}.")
