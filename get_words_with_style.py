import os
from docx import Document
import xml.etree.ElementTree as ET
from xml.dom import minidom
from docx.enum.text import WD_LINE_SPACING
import requests
from docx2python import docx2python

current_dir = os.path.dirname(os.path.realpath(__file__))

file_path = os.path.join(current_dir, 'test_docx.docx')

# Open Doc
doc = Document(file_path)


# для оффлайна сделать
def closest_color(requested_color):
    # Конвертируем RGB в HEX
    hex_value = '#{:02x}{:02x}{:02x}'.format(requested_color[0], requested_color[1], requested_color[2])
    response = requests.get(f"https://www.thecolorapi.com/id?hex={hex_value[1:]}")
    if response.status_code == 200:
        return response.json()['name']['value']
    else:
        return "Unknown"


def get_color_name(rgb_tuple):
    try:
        return closest_color(rgb_tuple)
    except Exception as e:
        print(f"Error converting color: {e}")
        return "Unknown"


# Get style text
def get_info_style(run, paragraph=None):
    font_name = run.font.name
    font_size = run.font.size.pt if run.font.size else None

    if not font_name and paragraph:
        font_name = paragraph.style.font.name if paragraph.style.font.name else "Unknown"
    if not font_size and paragraph:
        font_size = paragraph.style.font.size.pt if paragraph.style.font.size else "Unknown"

    bold = run.bold
    italic = run.italic

    # Get color font
    font_color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
    font_color_name = "Black"
    if font_color:
        try:
            font_color_name = get_color_name((font_color[0], font_color[1], font_color[2]))
        except ValueError:
            font_color_name = closest_color((font_color[0], font_color[1], font_color[2]))

    info_style = {
        'font': font_name,
        'size': f"{font_size}pt" if font_size != "Unknown" else "Unknown",
        'bold': 'yes' if bold else 'no',
        'italic': 'yes' if italic else 'no',
        'color': font_color_name
    }

    # Check result about style text
    print(
        f"Font: {info_style['font']} | Size: {info_style['size']} | Bold: {info_style['bold']} | Italic: {info_style['italic']} | Color: {info_style['color']}")

    return info_style


# Функция для получения информации о межстрочном интервале
def get_line_spacing_info(paragraph):
    para_format = paragraph.paragraph_format
    line_spacing = para_format.line_spacing
    line_spacing_rule = para_format.line_spacing_rule

    if line_spacing_rule == WD_LINE_SPACING.SINGLE:
        spacing_type = "Single"
    elif line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
        spacing_type = "1.5 lines"
    elif line_spacing_rule == WD_LINE_SPACING.DOUBLE:
        spacing_type = "Double"
    elif line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
        spacing_type = "At least"
    elif line_spacing_rule == WD_LINE_SPACING.EXACTLY:
        spacing_type = "Exactly"
    elif line_spacing_rule == WD_LINE_SPACING.MULTIPLE:
        spacing_type = "Multiple"
    else:
        spacing_type = "Unknown"

    # Print с информацией о межстрочном интервале
    print(f"Line Spacing Type: {spacing_type} | Line Spacing Value: {line_spacing}")

    return {
        'spacing_type': spacing_type,
        'line_spacing': line_spacing
    }


# Create XML element
def create_xml_element(element_type, text, info_style):
    element = ET.Element(element_type)
    element.text = text
    for key, value in info_style.items():
        if value:
            element.set(key, value)
    return element


# Функция для разбивки текста на слова, исключая пробелы и пустые строки
def split_text_into_words(text):
    return [word for word in text.split() if word.strip()]


# Функция для округления значений отступов и удаления нулей после запятой
def round_margin(value):
    rounded_value = round(value, 2)
    if rounded_value.is_integer():
        return int(rounded_value)
    return rounded_value


# Функция для получения параметров отступов страниц
def get_page_margins(doc):
    margins = []
    for section in doc.sections:
        margin_info = {
            'left_margin': round_margin(section.left_margin.cm),
            'right_margin': round_margin(section.right_margin.cm),
            'top_margin': round_margin(section.top_margin.cm),
            'bottom_margin': round_margin(section.bottom_margin.cm)
        }
        margins.append(margin_info)
        print(f"Left Margin: {margin_info['left_margin']} cm")
        print(f"Right Margin: {margin_info['right_margin']} cm")
        print(f"Top Margin: {margin_info['top_margin']} cm")
        print(f"Bottom Margin: {margin_info['bottom_margin']} cm")
    return margins


# Функция для получения сносок и их стилей
def get_footnotes_with_styles(docx_data):
    footnotes = []
    for i, footnote in enumerate(docx_data.footnotes):
        footnote_text = ' '.join([' '.join(cell) for row in footnote for cell in row])
        footnotes.append({'id': str(i + 1), 'text': footnote_text})
    return footnotes


# Функция для определения перекрестных ссылок
"""def is_cross_reference(run):
    if run.field and run.field.type == 'HYPERLINK':
        return True
    return False
"""

# Создание корневого элемента для нового XML
root = ET.Element("document")

# Получаем параметры отступов страниц и добавляем их в XML
page_margins = get_page_margins(doc)
for margin_info in page_margins:
    parameters_page_element = ET.Element("parameters_page")
    parameters_page_element.set("left_margin", str(margin_info['left_margin']))
    parameters_page_element.set("right_margin", str(margin_info['right_margin']))
    parameters_page_element.set("top_margin", str(margin_info['top_margin']))
    parameters_page_element.set("bottom_margin", str(margin_info['bottom_margin']))
    root.append(parameters_page_element)

# Проходим по всем параграфам в документе
for para in doc.paragraphs:
    if not para.text.strip():
        continue

    paragraph_element = ET.Element("paragraph")

    # Добавляем информацию о межстрочном интервале
    line_spacing_info = get_line_spacing_info(para)
    paragraph_element.set("line_spacing_type", line_spacing_info['spacing_type'])
    paragraph_element.set("line_spacing_value",
                          str(line_spacing_info['line_spacing']) if line_spacing_info['line_spacing'] else "Unknown")

    for run in para.runs:
        words = split_text_into_words(run.text)
        if not words:
            continue
        style_info = get_info_style(run, para)
        """for word in words:
            if is_cross_reference(run):
                cross_ref_element = create_xml_element("cross_reference", word, style_info)
                paragraph_element.append(cross_ref_element)
            else:
                word_element = create_xml_element("paragraph_words", word, style_info)
                paragraph_element.append(word_element)"""
    root.append(paragraph_element)

# Используем docx2python для извлечения сносок
docx_data = docx2python(file_path)
footnotes = get_footnotes_with_styles(docx_data)

# Добавляем сноски в XML
for footnote in footnotes:
    footnote_id = footnote['id']
    footnote_text = footnote['text']

    # Разделяем текст сносок на отдельные сноски
    footnote_parts = footnote_text.split('footnote')
    for part in footnote_parts:
        if part.strip():
            part_id, part_text = part.split(')', 1)
            footnote_element = ET.Element("footnote")
            footnote_element.set("id", part_id.strip())

            """# Создаем подэлемент с информацией о стиле
            sub_element = ET.Element(f"footnote{part_id.strip()}")
            sub_element.text = part_text.strip()

            # Добавляем информацию о стиле текста
            for run in doc.paragraphs[int(footnote_id) - 1].runs:
                style_info = get_info_style(run)
                for key, value in style_info.items():
                    sub_element.set(key, value)

            footnote_element.append(sub_element)"""
            root.append(footnote_element)
            # Print с информацией о сносках
            print(f"Footnote ID: {part_id.strip()} | Footnote Text: {part_text.strip()}")

# Сохраняем XML-документ в файл с отступами
output_file = os.path.join(current_dir, 'output_with_footnotes.xml')
tree = ET.ElementTree(root)

# Форматируем XML с отступами
xml_str = ET.tostring(root, encoding='utf-8')
parsed_str = minidom.parseString(xml_str)
pretty_xml_as_str = parsed_str.toprettyxml(indent="  ")

with open(output_file, 'w', encoding='utf-8') as f:
    f.write(pretty_xml_as_str)

print(f"Документ успешно сконвертирован в XML и сохранен как {output_file}.")
