import os
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
import lxml.etree as ET

current_dir = os.path.dirname(os.path.realpath(__file__))

file_path = os.path.join(current_dir, 'new_docx.docx')
xml_path = os.path.join(current_dir, 'output_ver2.xml')
output_file = os.path.join(current_dir, 'highlighted_docx.docx')


def get_style_info(run):
    font = run.font.name if run.font.name else "None"
    size = f"{run.font.size.pt} pt" if run.font.size else "None"
    bold = "Да" if run.bold else "Нет"
    return font, size, bold


def parse_paragraph(paragraph, is_toc=False):
    para_xml = ET.Element("toc_content" if is_toc else "paragraph")
    for run in paragraph.runs:
        font, size, bold = get_style_info(run)
        run_xml = ET.SubElement(para_xml, "content")
        run_xml.text = run.text
        run_xml.set("font", font)
        run_xml.set("size", size)
        run_xml.set("bold", bold)
    return para_xml


def parse_table(table):
    table_xml = ET.Element("table")
    for row in table.rows:
        row_xml = ET.SubElement(table_xml, "row")
        for cell in row.cells:
            cell_xml = ET.SubElement(row_xml, "cell")
            for paragraph in cell.paragraphs:
                cell_xml.append(parse_paragraph(paragraph))
    return table_xml


def parse_document(doc):
    root = ET.Element("document")
    is_toc_section = False

    for paragraph in doc.paragraphs:
        # Проверка на заголовок оглавления
        if paragraph.text.strip().lower() in ["оглавление", "содержание"]:
            is_toc_section = True
            continue

        # Если это оглавление, обрабатываем как toc_content
        if is_toc_section:
            root.append(parse_paragraph(paragraph, is_toc=True))
            is_toc_section = False
        else:
            root.append(parse_paragraph(paragraph))

    for table in doc.tables:
        root.append(parse_table(table))

    return root


def save_to_xml(root, filename):
    tree = ET.ElementTree(root)
    with open(filename, 'wb') as f:
        tree.write(f, encoding="utf-8", xml_declaration=True, pretty_print=True)


def highlight_text(run):
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_comment(paragraph, comment_text):
    # Добавляем комментарий к параграфу
    paragraph.add_comment(comment_text, author='System')


def process_document(doc, font_info):
    for para_info, paragraph in zip(font_info, doc.paragraphs):
        for (text, font, size, bold), run in zip(para_info, paragraph.runs):
            if font != "Times New Roman":
                highlight_text(run)
                add_comment(paragraph, "Выберите корректный шрифт")


def parse_xml(xml_filename):
    if not os.path.exists(xml_filename):
        raise FileNotFoundError(f"XML file '{xml_filename}' not found.")

    tree = ET.parse(xml_filename)
    root = tree.getroot()
    font_info = []

    for paragraph in root.findall('paragraph'):
        para_info = []
        for content in paragraph.findall('content'):
            text = content.text
            font = content.get('font')
            size = content.get('size')
            bold = content.get('bold')
            para_info.append((text, font, size, bold))
        font_info.append(para_info)

    return font_info


def main(docx_filename, xml_filename, output_filename):
    doc = docx.Document(docx_filename)
    font_info = parse_xml(xml_filename)
    process_document(doc, font_info)
    doc.save(output_filename)


if __name__ == "__main__":
    main(file_path, xml_path, output_file)